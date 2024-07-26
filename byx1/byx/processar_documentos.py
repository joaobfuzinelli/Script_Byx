import os
import openpyxl
from docx import Document
from docx.shared import Pt
from datetime import datetime
import locale

def ler_excel(nome_arquivo):
    wb = openpyxl.load_workbook(nome_arquivo)
    planilha = wb.active

    dados = []

    for row in planilha.iter_rows(min_row=2, values_only=True): 
        nome, cpf, contrato, status, obs = row[2], row[3], row[6], row[9], row[8]
        if not status:  # Verifica se a coluna Status está vazia
            dados.append({'Nome': nome, 'CPF': cpf, 'Contrato': contrato, 'Obs': obs})

    return dados

def formatar_data():
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
    return datetime.now().strftime("São Paulo, %d de %B de %Y")

def aplicar_formatacao(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Verdana'
            run.font.size = Pt(9)
    return doc

def preencher_template_word(dados, templates):
    data_atual = formatar_data()  # Formata a data no formato desejado
    pasta_destino = f"arquivos_processados_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}"
    
    if not os.path.exists(pasta_destino):
        os.makedirs(pasta_destino)

    categorias_validas = ['RETIDO', 'PORTADO', 'SEM PEDIDO']
    
    for item in dados:
        status = item['Obs']
        
        # Verifica se o status é válido e cria a subpasta se necessário
        if status in categorias_validas:
            pasta_status = os.path.join(pasta_destino, status)
            if not os.path.exists(pasta_status):
                os.makedirs(pasta_status)
        else:
            pasta_status = pasta_destino  # Usa a pasta principal para status desconhecidos
        
        template_doc = templates.get(status, templates.get('OUTROS'))  # Usa o template padrão se necessário
        if not template_doc:
            print(f"Template não encontrado para Obs: {status}")
            continue
        
        doc = Document(template_doc)  # Carrega o template a cada iteração para iniciar limpo

        for paragraph in doc.paragraphs:
            if '<NOME>' in paragraph.text and item['Nome']:
                paragraph.text = paragraph.text.replace('<NOME>', item['Nome'])
            if '<CPF>' in paragraph.text and item['CPF']:
                paragraph.text = paragraph.text.replace('<CPF>', item['CPF'])
            if '<CONTRATO>' in paragraph.text and item['Contrato']:
                paragraph.text = paragraph.text.replace('<CONTRATO>', item['Contrato'])
            if '<DATA>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<DATA>', data_atual)

        doc = aplicar_formatacao(doc)  # Aplica a formatação Verdana 9

        nome_arquivo = f"{item['Nome']}.docx"  # Define o nome do arquivo com base no nome preenchido
        caminho_arquivo = os.path.join(pasta_status, nome_arquivo)  # Caminho completo para salvar o arquivo
        doc.save(caminho_arquivo)
        print(f'Documento preenchido e salvo como {caminho_arquivo}')

# Caminho do arquivo Excel
excel_file = "PROCON_BACEN.xlsx"

# Caminhos dos templates do documento Word
templates = {
    'PORTADO': 'PORTADO.docx',
    'RETIDO': 'RETIDO.docx',
    'SEM PEDIDO': 'SEM_PEDIDO.docx'
}

# Ler apenas as colunas Nome, CPF, Contrato, Status e Obs do Excel
dados = ler_excel(excel_file)

# Preencher o documento Word com os dados lidos do Excel e salvar na pasta correspondente
preencher_template_word(dados, templates)


input("Aperte uma tecla para fechar: ")
